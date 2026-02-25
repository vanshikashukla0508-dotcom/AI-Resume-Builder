import streamlit as st
import os
from openai import OpenAI
from docx import Document
import io
# -------------------------
# DOCX RESUME CREATOR FUNCTION
# -------------------------
def create_docx_resume(resume_text, template):

    doc = Document()

    if template == "Modern":
        doc.add_heading("RESUME", level=0)

    elif template == "Professional":
        doc.add_heading("Professional Resume", level=0)

    elif template == "Minimal":
        doc.add_heading("", level=0)

    else:
        doc.add_heading("Resume", level=0)

    for line in resume_text.split("\n"):
        doc.add_paragraph(line)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer
# -----------------------------
# CONFIGURATION
# -----------------------------
import streamlit as st
from openai import OpenAI

# Initialize the OpenAI client correctly for a real OpenAI Key
client = OpenAI(
    api_key=st.secrets["OPENAI_API_KEY"]
)


MODEL_NAME = "gpt-3.5-turbo"
st.set_page_config(page_title="AI Resume Builder", layout="wide")

st.title("🚀 AI Resume Builder Pro")
st.caption("Created by Vanshika Shukla")
st.write("Generate Professional, ATS-Optimized Resume with AI")

# -----------------------------
# TEMPLATE SELECTION
# -----------------------------
template = st.selectbox(
    "Choose Resume Template",
    ["Professional", "Modern", "Minimal", "Creative"]
)

st.markdown("---")

# -----------------------------
# USER INPUT SECTION
# -----------------------------
name = st.text_input("Full Name")
contact = st.text_input("Contact Information (Phone | Email | Address)")

education = st.text_area("Education")
skills = st.text_area("Skills")
experience = st.text_area("Experience")
projects = st.text_area("Projects")
certifications = st.text_area("Certifications")
achievements = st.text_area("Achievements")
positions = st.text_area("Positions of Responsibility")
extra_curricular = st.text_area("Extra-Curricular Activities")
career_goal = st.text_area("Career Objective")

st.markdown("---")

# -------------------------
# GENERATE RESUME
# -------------------------

if st.button("Generate Full Resume"):

    if name and education and skills:

        prompt = f"""
        Create a {template} style professional ATS-friendly resume.

        Include these sections:
        1. Contact Information
        2. Professional Summary
        3. Education
        4. Skills
        5. Experience
        6. Projects
        7. Certifications
        8. Achievements
        9. Positions of Responsibility
        10. Extra-Curricular Activities
        11. Career Objective
        12. References

        USER INFORMATION:
        Name: {name}
        Contact: {contact}
        Education: {education}
        Skills: {skills}
        Experience: {experience}
        Projects: {projects}
        Certifications: {certifications}
        Achievements: {achievements}
        Positions: {positions}
        Extra Curricular: {extra_curricular}
        Career Objective: {career_goal}

        Make it highly professional and impactful.
        """

        response = client.chat.completions.create(
            model=MODEL_NAME,
            messages=[{"role": "user", "content": prompt}]
        )

        resume_output = response.choices[0].message.content

        st.subheader("📄 Generated Resume")
        st.write(resume_output)

        # Create DOCX
        docx_file = create_docx_resume(resume_output, template)

        st.download_button(
            label="Download Resume (.docx)",
            data=docx_file,
            file_name=f"{name}_Resume_{template}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    else:
        st.warning("Please fill at least Name, Education and Skills")

# -----------------------------
# LINKEDIN SUMMARY GENERATOR
# -----------------------------
if st.button("Generate LinkedIn Summary"):

    linkedin_prompt = f"""
    Create a powerful LinkedIn summary based on:

    Name: {name}
    Education: {education}
    Skills: {skills}
    Experience: {experience}
    Career Goal: {career_goal}

    Make it professional, confident, and engaging.
    """

    linkedin_response = client.chat.completions.create(
        model=MODEL_NAME,
        messages=[{"role": "user", "content": linkedin_prompt}]
    )

    st.subheader("💼 LinkedIn Summary")
    st.write(linkedin_response.choices[0].message.content)

st.markdown("---")

# -----------------------------
# ATS SCORE CHECKER
# -----------------------------
if st.button("Check ATS Score"):

    ats_prompt = f"""
    Evaluate the following resume information for ATS compatibility.
    Give:
    - ATS Score out of 100
    - Strengths
    - Weaknesses
    - Improvement Suggestions

    Resume Data:
    Skills: {skills}
    Experience: {experience}
    Projects: {projects}
    Certifications: {certifications}
    Achievements: {achievements}
    """

    ats_response = client.chat.completions.create(
        model=MODEL_NAME,
        messages=[{"role": "user", "content": ats_prompt}]
    )

    st.subheader("📊 ATS Evaluation Report")
    st.write(ats_response.choices[0].message.content)

st.markdown("---")

# -----------------------------
# RESUME IMPROVEMENT SUGGESTIONS
# -----------------------------
if st.button("Suggest Improvements"):

    improve_prompt = f"""
    Suggest professional improvements for this resume profile:

    Education: {education}
    Skills: {skills}
    Experience: {experience}
    Projects: {projects}
    Career Goal: {career_goal}

    Give actionable suggestions.
    """

    improve_response = client.chat.completions.create(
        model=MODEL_NAME,
        messages=[{"role": "user", "content": improve_prompt}]
    )

    st.subheader("✨ Resume Improvement Suggestions")
    st.write(improve_response.choices[0].message.content)
    # ----------------------------
# DOCX RESUME GENERATOR
# ----------------------------

from docx import Document
from docx.shared import Pt
from io import BytesIO


def create_docx_resume(resume_text, template_style):

    doc = Document()

    # TEMPLATE STYLING
    if template_style == "Professional":
        font_name = "Calibri"
        heading_size = 16

    elif template_style == "Modern":
        font_name = "Arial"
        heading_size = 18

    elif template_style == "Minimal":
        font_name = "Times New Roman"
        heading_size = 14

    elif template_style == "Creative":
        font_name = "Georgia"
        heading_size = 20

    # Add Resume Content
    for line in resume_text.split("\n"):

        paragraph = doc.add_paragraph()
        run = paragraph.add_run(line)

        run.font.name = font_name
        run.font.size = Pt(11)

        if line.isupper() or ":" in line:
            run.bold = True
            run.font.size = Pt(heading_size)

    # Save to memory
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer
# ==========================
# FOOTER
# ==========================

st.markdown("---")
st.markdown(
    "<center>© 2026 Created by <b>Vanshika Shukla</b> | AI Resume Builder Project</center>",
    unsafe_allow_html=True
)